from pathlib import Path
import re

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn


ROOT = Path(__file__).resolve().parent
MD_PATH = ROOT / "bao_cao_qua_trinh_train_models.md"
DOCX_PATH = ROOT / "bao_cao_qua_trinh_train_models.docx"


def shade_cell(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for name, value in [("top", top), ("start", start), ("bottom", bottom), ("end", end)]:
        node = tc_mar.find(qn(f"w:{name}"))
        if node is None:
            node = OxmlElement(f"w:{name}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_borders(table, color="DADCE0", sz="4"):
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.first_child_found_in("w:tblBorders")
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = f"w:{edge}"
        element = borders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), sz)
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), color)


def repeat_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def add_inline_runs(paragraph, text):
    # Minimal Markdown inline support for backticks and bold markers.
    token_re = re.compile(r"(`[^`]+`|\*\*[^*]+\*\*)")
    pos = 0
    for match in token_re.finditer(text):
        if match.start() > pos:
            paragraph.add_run(text[pos:match.start()])
        token = match.group(0)
        run = paragraph.add_run(token[1:-1] if token.startswith("`") else token[2:-2])
        if token.startswith("`"):
            run.font.name = "Consolas"
            run._element.rPr.rFonts.set(qn("w:eastAsia"), "Consolas")
            run.font.size = Pt(9.5)
            run.font.color.rgb = RGBColor(80, 80, 80)
        else:
            run.bold = True
        pos = match.end()
    if pos < len(text):
        paragraph.add_run(text[pos:])


def parse_table(lines, start):
    rows = []
    idx = start
    while idx < len(lines) and lines[idx].strip().startswith("|"):
        parts = [p.strip() for p in lines[idx].strip().strip("|").split("|")]
        rows.append(parts)
        idx += 1
    if len(rows) >= 2 and all(set(c) <= {"-", ":"} for c in rows[1]):
        return rows[0], rows[2:], idx
    return None, None, start


def add_table(doc, headers, rows):
    cols = len(headers)
    table = doc.add_table(rows=1, cols=cols)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    set_table_borders(table)
    font_size = 7 if cols >= 5 else 8.5
    widths = [6.5 / cols] * cols
    if cols == 3:
        widths = [1.7, 1.7, 3.1]
    elif cols == 6:
        widths = [0.9, 1.15, 0.85, 1.45, 1.0, 1.15]

    repeat_header(table.rows[0])
    for i, header in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.width = Inches(widths[i])
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        set_cell_margins(cell)
        shade_cell(cell, "F2F4F7")
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(header)
        run.bold = True
        run.font.size = Pt(font_size)

    for row in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row[:cols]):
            cell = cells[i]
            cell.width = Inches(widths[i])
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            set_cell_margins(cell)
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            add_inline_runs(p, value)
            for run in p.runs:
                if run.font.size is None or run.font.size.pt > font_size:
                    run.font.size = Pt(font_size)
    doc.add_paragraph()


def add_callout(doc):
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_borders(table, color="B7C9E2", sz="6")
    cell = table.cell(0, 0)
    shade_cell(cell, "F4F6F9")
    set_cell_margins(cell, top=140, bottom=140, start=180, end=180)
    p = cell.paragraphs[0]
    r = p.add_run("Kết luận nhanh: ")
    r.bold = True
    r.font.color.rgb = RGBColor(31, 78, 121)
    p.add_run(
        "ứng viên phân loại thuyết phục nhất là VGG16+Inception+ViT v4; "
        "ConvNeXtV2 là lựa chọn gọn, mạnh và dễ bảo trì; YOLOv8s là nhánh detection lesion tốt nhất theo output hiện có."
    )
    doc.add_paragraph()


def build():
    text = MD_PATH.read_text(encoding="utf-8")
    lines = text.splitlines()
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)

    styles = doc.styles
    styles["Normal"].font.name = "Calibri"
    styles["Normal"]._element.rPr.rFonts.set(qn("w:eastAsia"), "Calibri")
    styles["Normal"].font.size = Pt(11)
    styles["Normal"].paragraph_format.space_after = Pt(6)
    styles["Normal"].paragraph_format.line_spacing = 1.10
    for name, size, color in [
        ("Heading 1", 16, RGBColor(31, 78, 121)),
        ("Heading 2", 13, RGBColor(31, 78, 121)),
        ("Heading 3", 12, RGBColor(31, 77, 120)),
    ]:
        style = styles[name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Calibri")
        style.font.size = Pt(size)
        style.font.color.rgb = color
        style.font.bold = True
        style.paragraph_format.space_before = Pt(12 if name != "Heading 1" else 16)
        style.paragraph_format.space_after = Pt(6)

    title = lines[0].lstrip("# ").strip()
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run(title)
    run.bold = True
    run.font.size = Pt(22)
    run.font.color.rgb = RGBColor(11, 37, 69)
    sub = doc.add_paragraph()
    sub.paragraph_format.space_after = Pt(14)
    sub_run = sub.add_run(
        "Phân loại dựa trên code trong thư mục train files, kèm so sánh ưu điểm và khuyến nghị chọn mô hình"
    )
    sub_run.font.color.rgb = RGBColor(89, 89, 89)
    add_callout(doc)

    i = 1
    while i < len(lines):
        line = lines[i].strip()
        if not line:
            i += 1
            continue
        if line.startswith("|"):
            headers, rows, next_i = parse_table(lines, i)
            if headers:
                add_table(doc, headers, rows)
                i = next_i
                continue
        if line.startswith("### "):
            doc.add_heading(line[4:].strip(), level=3)
        elif line.startswith("## "):
            doc.add_heading(line[3:].strip(), level=1)
        elif line.startswith("# "):
            pass
        elif line.startswith("- "):
            p = doc.add_paragraph(style="List Bullet")
            add_inline_runs(p, line[2:].strip())
        elif re.match(r"^\d+\.\s+", line):
            p = doc.add_paragraph(style="List Number")
            add_inline_runs(p, re.sub(r"^\d+\.\s+", "", line))
        else:
            p = doc.add_paragraph()
            add_inline_runs(p, line)
        i += 1

    for sec in doc.sections:
        footer = sec.footer.paragraphs[0]
        footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        run = footer.add_run("Báo cáo train models - tạo từ phân tích code notebook")
        run.font.size = Pt(8)
        run.font.color.rgb = RGBColor(89, 89, 89)

    doc.save(DOCX_PATH)
    print(DOCX_PATH)


if __name__ == "__main__":
    build()

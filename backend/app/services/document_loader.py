from __future__ import annotations

import logging
import re
from dataclasses import dataclass
from pathlib import Path

logger = logging.getLogger(__name__)

SUPPORTED_EXTENSIONS = {".md", ".txt", ".pdf", ".docx"}


@dataclass(frozen=True)
class LoadedDocument:
    text: str
    source: str
    file_path: str
    page: int | None = None


def load_documents(
    knowledge_base_path: str | Path,
    docs_dir: str | Path | None = None,
) -> list[LoadedDocument]:
    """Load Markdown, TXT, PDF and DOCX sources into text documents."""
    docs: list[LoadedDocument] = []
    kb_path = Path(knowledge_base_path)

    if kb_path.exists():
        docs.extend(load_document(kb_path))

    if docs_dir:
        root = Path(docs_dir)
        if root.exists():
            for path in sorted(root.rglob("*")):
                if not path.is_file() or path.suffix.lower() not in SUPPORTED_EXTENSIONS:
                    continue
                if kb_path.exists() and path.resolve() == kb_path.resolve():
                    continue
                docs.extend(load_document(path))

    return [doc for doc in docs if doc.text.strip()]


def load_document(path: str | Path) -> list[LoadedDocument]:
    path = Path(path)
    suffix = path.suffix.lower()

    if suffix in {".md", ".txt"}:
        return [_load_text(path)]
    if suffix == ".pdf":
        return _load_pdf(path)
    if suffix == ".docx":
        return [_load_docx(path)]
    return []


def _load_text(path: Path) -> LoadedDocument:
    try:
        text = path.read_text(encoding="utf-8")
    except UnicodeDecodeError:
        text = path.read_text(encoding="utf-8-sig", errors="ignore")
    return LoadedDocument(text=text, source=path.name, file_path=str(path))


def _load_pdf(path: Path) -> list[LoadedDocument]:
    """Load PDF with improved text extraction.

    Strategy:
    1. Try pdfplumber first (better layout/table handling).
    2. Fall back to pypdf if pdfplumber is not installed.
    3. Post-process extracted text to fix common PDF artifacts.
    """
    try:
        return _load_pdf_pdfplumber(path)
    except ImportError:
        pass

    try:
        return _load_pdf_pypdf(path)
    except ImportError as exc:
        raise ImportError(
            "Install pypdf or pdfplumber to load PDF knowledge files: "
            "pip install pypdf  OR  pip install pdfplumber"
        ) from exc


def _load_pdf_pdfplumber(path: Path) -> list[LoadedDocument]:
    import pdfplumber  # type: ignore[import]

    docs: list[LoadedDocument] = []
    with pdfplumber.open(str(path)) as pdf:
        for index, page in enumerate(pdf.pages, start=1):
            # Extract tables first, then remaining text
            parts: list[str] = []

            tables = page.extract_tables()
            table_bboxes = [t.bbox for t in page.find_tables()] if tables else []

            # Text outside tables
            if table_bboxes:
                text_outside = page.filter(
                    lambda obj, bboxes=table_bboxes: not any(
                        _bbox_overlap(obj["x0"], obj["top"], obj["x1"], obj["bottom"], b)
                        for b in bboxes
                    )
                ).extract_text()
            else:
                text_outside = page.extract_text()

            if text_outside and text_outside.strip():
                parts.append(_clean_pdf_text(text_outside))

            # Tables as pipe-separated rows
            for table in tables:
                rows = [
                    " | ".join(str(cell or "").strip() for cell in row)
                    for row in table
                    if any(cell and str(cell).strip() for cell in row)
                ]
                if rows:
                    parts.append("\n".join(rows))

            combined = "\n\n".join(p for p in parts if p.strip())
            if combined.strip():
                docs.append(
                    LoadedDocument(
                        text=combined,
                        source=path.name,
                        file_path=str(path),
                        page=index,
                    )
                )
    return docs


def _load_pdf_pypdf(path: Path) -> list[LoadedDocument]:
    from pypdf import PdfReader  # type: ignore[import]

    reader = PdfReader(str(path))
    docs: list[LoadedDocument] = []
    for index, page in enumerate(reader.pages, start=1):
        text = page.extract_text() or ""
        cleaned = _clean_pdf_text(text)
        if cleaned.strip():
            docs.append(
                LoadedDocument(
                    text=cleaned,
                    source=path.name,
                    file_path=str(path),
                    page=index,
                )
            )
    return docs


def _load_docx(path: Path) -> LoadedDocument:
    try:
        import docx  # type: ignore[import]
    except ImportError as exc:
        raise ImportError(
            "Install python-docx to load DOCX knowledge files: pip install python-docx"
        ) from exc

    document = docx.Document(str(path))
    parts: list[str] = []

    for paragraph in document.paragraphs:
        stripped = paragraph.text.strip()
        if not stripped:
            continue
        # Preserve heading structure so chunker can detect sections
        style_name = (paragraph.style.name or "").lower()
        if "heading" in style_name:
            # Convert DOCX headings to Markdown-style so _split_sections picks them up
            try:
                level = int(re.search(r"\d+", style_name).group())  # type: ignore[union-attr]
            except (AttributeError, ValueError):
                level = 1
            parts.append(f"{'#' * level} {stripped}")
        else:
            parts.append(stripped)

    for table in document.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if cells:
                parts.append(" | ".join(cells))

    return LoadedDocument(text="\n\n".join(parts), source=path.name, file_path=str(path))


# ── PDF text post-processing ─────────────────────────────────────────

def _clean_pdf_text(text: str) -> str:
    """Fix common PDF extraction artifacts."""
    if not text:
        return text

    # Normalize line endings
    text = re.sub(r"\r\n?", "\n", text)

    # Remove hyphenation at line breaks (e.g. "treat-\nment" → "treatment")
    text = re.sub(r"(\w)-\n(\w)", r"\1\2", text)

    # Collapse lines that are clearly continuation (no sentence-ending punctuation)
    # but keep paragraph breaks (double newlines)
    lines = text.split("\n")
    merged: list[str] = []
    for line in lines:
        stripped = line.strip()
        if not stripped:
            merged.append("")
            continue
        if (
            merged
            and merged[-1]
            and not merged[-1].endswith((".", "!", "?", ":", ";"))
            and not stripped[0].isupper()
        ):
            merged[-1] = merged[-1] + " " + stripped
        else:
            merged.append(stripped)

    # Collapse multiple blank lines into one
    result = re.sub(r"\n{3,}", "\n\n", "\n".join(merged))
    return result.strip()


def _bbox_overlap(x0: float, top: float, x1: float, bottom: float, bbox: tuple) -> bool:
    """Check if a character bounding box overlaps with a table bounding box."""
    bx0, btop, bx1, bbottom = bbox
    return not (x1 < bx0 or x0 > bx1 or bottom < btop or top > bbottom)

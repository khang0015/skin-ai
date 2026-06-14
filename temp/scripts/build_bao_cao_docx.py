from __future__ import annotations

import re
import shutil
import sys
from datetime import datetime
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt


ROOT = Path(__file__).resolve().parents[1]
MD_PATH = ROOT / "bao_cao_so_bo.md"
DOCX_PATH = ROOT / "bao_cao_chinh_sua.docx"
FALLBACK_DOCX_PATH = ROOT / "bao_cao_chinh_sua_moi.docx"


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=90, start=90, bottom=90, end=90) -> None:
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for m, v in {"top": top, "start": start, "bottom": bottom, "end": end}.items():
        node = tc_mar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(v))
        node.set(qn("w:type"), "dxa")


def add_page_number(paragraph) -> None:
    run = paragraph.add_run()
    fld_char1 = OxmlElement("w:fldChar")
    fld_char1.set(qn("w:fldCharType"), "begin")
    instr_text = OxmlElement("w:instrText")
    instr_text.set(qn("xml:space"), "preserve")
    instr_text.text = "PAGE"
    fld_char2 = OxmlElement("w:fldChar")
    fld_char2.set(qn("w:fldCharType"), "end")
    run._r.append(fld_char1)
    run._r.append(instr_text)
    run._r.append(fld_char2)


def clean_inline(text: str) -> str:
    text = text.strip()
    text = re.sub(r"`([^`]+)`", r"\1", text)
    text = text.replace("**", "")
    text = text.replace("*", "")
    text = text.replace("“", '"').replace("”", '"')
    return text


def is_table_start(lines: list[str], idx: int) -> bool:
    if idx + 1 >= len(lines):
        return False
    return lines[idx].lstrip().startswith("|") and re.match(r"^\s*\|?\s*:?-{3,}:?\s*(\|\s*:?-{3,}:?\s*)+\|?\s*$", lines[idx + 1])


def parse_table(lines: list[str], idx: int) -> tuple[list[list[str]], int]:
    rows: list[list[str]] = []
    while idx < len(lines) and lines[idx].lstrip().startswith("|"):
        line = lines[idx].strip()
        if re.match(r"^\|?\s*:?-{3,}:?\s*(\|\s*:?-{3,}:?\s*)+\|?$", line):
            idx += 1
            continue
        cells = [clean_inline(c.strip()) for c in line.strip("|").split("|")]
        rows.append(cells)
        idx += 1
    return rows, idx


def set_normal_style(doc: Document) -> None:
    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Times New Roman"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    normal.font.size = Pt(13)
    pf = normal.paragraph_format
    pf.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    pf.line_spacing = 1.3
    pf.space_before = Pt(0)
    pf.space_after = Pt(6)
    pf.first_line_indent = Cm(0.7)

    for name in ["Heading 1", "Heading 2", "Heading 3"]:
        st = styles[name]
        st.font.name = "Times New Roman"
        st._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
        st.font.color.rgb = None
        st.font.size = Pt(14 if name == "Heading 1" else 13)
        st.font.bold = True
        st.font.italic = name == "Heading 3"
        st.paragraph_format.line_spacing = 1.3
        st.paragraph_format.space_before = Pt(6 if name != "Heading 1" else 0)
        st.paragraph_format.space_after = Pt(6)
        st.paragraph_format.first_line_indent = None


def configure_sections(doc: Document) -> None:
    section = doc.sections[0]
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.left_margin = Cm(3)
    section.right_margin = Cm(2)
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.header_distance = Cm(1.6)
    section.footer_distance = Cm(1.6)

    header_p = section.header.paragraphs[0]
    header_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    header_p.text = "Xây dựng ứng dụng hỗ trợ chẩn đoán bệnh da liễu"
    for run in header_p.runs:
        run.font.name = "Times New Roman"
        run._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
        run.font.size = Pt(10)
        run.font.italic = True

    footer_p = section.footer.paragraphs[0]
    footer_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer_p.add_run("Sinh viên thực hiện: [Điền họ tên]    GVHD: [Điền họ tên]    Trang ")
    add_page_number(footer_p)
    for run in footer_p.runs:
        run.font.name = "Times New Roman"
        run._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
        run.font.size = Pt(10)


def add_heading(doc: Document, text: str, level: int) -> None:
    if level == 1 and (text.startswith("CHƯƠNG") or text in {"MỞ ĐẦU", "KẾT LUẬN", "TÀI LIỆU THAM KHẢO", "PHỤ LỤC A: GHI CHÚ CẦN BỔ SUNG CHO BẢN CHÍNH THỨC"}):
        if len(doc.paragraphs) > 1:
            doc.add_page_break()
    p = doc.add_paragraph(style=f"Heading {min(level, 3)}")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER if level == 1 else WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.first_line_indent = None
    run = p.add_run(clean_inline(text))
    run.font.name = "Times New Roman"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    run.font.size = Pt(14 if level == 1 else 13)
    run.bold = True
    run.italic = level == 3


def add_body_paragraph(doc: Document, text: str, *, italic=False, center=False, no_indent=False) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER if center else WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.line_spacing = 1.3
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.first_line_indent = None if no_indent or center else Cm(0.7)
    run = p.add_run(clean_inline(text))
    run.font.name = "Times New Roman"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    run.font.size = Pt(13)
    run.italic = italic


def add_table(doc: Document, rows: list[list[str]]) -> None:
    if not rows:
        return
    cols = max(len(r) for r in rows)
    table = doc.add_table(rows=len(rows), cols=cols)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    for r_idx, row in enumerate(rows):
        for c_idx in range(cols):
            cell = table.cell(r_idx, c_idx)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            set_cell_margins(cell)
            if c_idx < len(row):
                cell.text = row[c_idx]
            else:
                cell.text = ""
            for p in cell.paragraphs:
                p.paragraph_format.line_spacing = 1.15
                p.paragraph_format.space_after = Pt(0)
                p.paragraph_format.first_line_indent = None
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER if r_idx == 0 or len(cell.text) < 24 else WD_ALIGN_PARAGRAPH.JUSTIFY
                for run in p.runs:
                    run.font.name = "Times New Roman"
                    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
                    run.font.size = Pt(12)
                    run.bold = r_idx == 0
            if r_idx == 0:
                set_cell_shading(cell, "D9EAF7")
    doc.add_paragraph()


def resolve_root_path(path_text: str) -> Path:
    path = Path(path_text)
    if path.is_absolute():
        return path
    return ROOT / path


def build_docx(md_path: Path = MD_PATH, docx_path: Path = DOCX_PATH) -> None:
    md = md_path.read_text(encoding="utf-8")
    lines = md.splitlines()
    doc = Document()
    set_normal_style(doc)
    configure_sections(doc)

    in_fence = False
    fence_type = ""
    i = 0
    while i < len(lines):
        raw = lines[i]
        line = raw.rstrip()
        stripped = line.strip()

        if stripped.startswith("```"):
            if not in_fence:
                in_fence = True
                fence_type = stripped.removeprefix("```").strip()
            else:
                in_fence = False
                fence_type = ""
            i += 1
            continue

        if in_fence:
            if fence_type == "text" and stripped and not stripped.startswith("->"):
                add_body_paragraph(doc, stripped, no_indent=True)
            i += 1
            continue

        if not stripped:
            i += 1
            continue

        if is_table_start(lines, i):
            rows, i = parse_table(lines, i)
            add_table(doc, rows)
            continue

        if stripped.startswith("# "):
            add_heading(doc, stripped[2:].strip(), 1)
        elif stripped.startswith("## "):
            add_heading(doc, stripped[3:].strip(), 2)
        elif stripped.startswith("### "):
            add_heading(doc, stripped[4:].strip(), 3)
        elif stripped.startswith(">"):
            text = stripped.lstrip(">").strip()
            add_body_paragraph(doc, text, italic=True, center=text.startswith("Chèn") or text.startswith("Ghi chú"), no_indent=True)
        elif re.match(r"^(Hình|Bảng) \d+\.\d+", stripped):
            add_body_paragraph(doc, stripped, center=True, no_indent=True)
        elif stripped.startswith("- "):
            p = doc.add_paragraph(style=None)
            p.paragraph_format.left_indent = Cm(0.7)
            p.paragraph_format.first_line_indent = Cm(-0.3)
            p.paragraph_format.line_spacing = 1.3
            run = p.add_run("• " + clean_inline(stripped[2:]))
            run.font.name = "Times New Roman"
            run._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
            run.font.size = Pt(13)
        else:
            add_body_paragraph(doc, stripped)
        i += 1

    if docx_path.exists():
        backup = docx_path.with_name(f"{docx_path.stem}.backup_{datetime.now():%Y%m%d_%H%M%S}.docx")
        shutil.copy2(docx_path, backup)
        print(f"Backup: {backup}")
    try:
        doc.save(docx_path)
        print(f"Wrote: {docx_path}")
    except PermissionError:
        try:
            fallback = docx_path.with_name(f"{docx_path.stem}_moi{docx_path.suffix}")
            doc.save(fallback)
            print(f"Wrote fallback because {docx_path.name} is locked: {fallback}")
        except PermissionError:
            versioned = docx_path.with_name(f"{docx_path.stem}_{datetime.now():%Y%m%d_%H%M%S}{docx_path.suffix}")
            doc.save(versioned)
            print(f"Wrote versioned fallback because both DOCX files are locked: {versioned}")


if __name__ == "__main__":
    input_path = resolve_root_path(sys.argv[1]) if len(sys.argv) > 1 else MD_PATH
    output_path = resolve_root_path(sys.argv[2]) if len(sys.argv) > 2 else DOCX_PATH
    build_docx(input_path, output_path)
